#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
References.py
Kunio Tsujita articles' builder - Updated with external styles module and code support
"""

import os
import re
import sys
import argparse
import subprocess
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

# Import external styles module
from DocxStylesNewer import create_styles_for_document, StylesManager, modify_styles

try:
    from pypdf import PdfReader
except ImportError:
    print("ERROR: pypdf not installed. Install with: pip install pypdf")
    sys.exit(1)


class MarkdownDocxParser:
    """Parser for converting markdown to DOCX elements with formatting preservation"""

    def __init__(self, document, styles_manager):
        self.document = document
        self.styles_manager = styles_manager
        self.in_code_block = False
        self.code_block_lines = []
        self.code_language = ""

    def detect_japanese(self, text):
        """Detect if text contains Japanese characters"""
        japanese_pattern = r'[\u3040-\u309F\u30A0-\u30FF\u4E00-\u9FAF]+'
        return bool(re.search(japanese_pattern, text))

    def process_formatted_text(self, text, paragraph):
        """Process markdown formatting and add runs to paragraph with proper Japanese handling"""

        def split_text_by_language(text_segment):
            """Split text into Japanese and non-Japanese parts"""
            parts = []
            current_part = ""
            current_is_japanese = False

            for char in text_segment:
                char_is_japanese = bool(re.match(r'[\u3040-\u309F\u30A0-\u30FF\u4E00-\u9FAF]', char))

                if char_is_japanese != current_is_japanese:
                    # Language changed, save current part
                    if current_part:
                        parts.append({
                            'text': current_part,
                            'is_japanese': current_is_japanese
                        })
                    current_part = char
                    current_is_japanese = char_is_japanese
                else:
                    current_part += char

            # Add final part
            if current_part:
                parts.append({
                    'text': current_part,
                    'is_japanese': current_is_japanese
                })

            return parts

        def add_text_run(text_content, bold=False, italic=False, code=False):
            """Add text run with proper font handling"""
            if not text_content:
                return

            if code:
                # For inline code, use monospace font regardless of language
                run = paragraph.add_run(text_content)
                run.font.name = self.styles_manager.document_styles.font_manager.get_font('code')
                run.font.size = Pt(self.styles_manager.config.font_sizes['code_inline'])
                run.font.color.rgb = self.styles_manager.config.colors['code_text']
                return

            # Split by language and create separate runs
            language_parts = split_text_by_language(text_content)

            for part in language_parts:
                run = paragraph.add_run(part['text'])

                # Apply formatting
                if bold:
                    run.bold = True
                if italic:
                    run.italic = True

                # Apply appropriate font
                if part['is_japanese']:
                    run.font.name = self.styles_manager.document_styles.font_manager.get_font('japanese')

        def parse_nested_formatting(text_segment):
            """Parse nested formatting patterns in text segment"""
            # First, handle inline code to avoid conflicts with other formatting
            inline_code_pattern = r'`([^`]+)`'
            inline_code_segments = []

            for match in re.finditer(inline_code_pattern, text_segment):
                inline_code_segments.append({
                    'start': match.start(),
                    'end': match.end(),
                    'content': match.group(1),
                    'type': 'inline_code',
                    'original': match.group(0)
                })

            # Replace inline code with placeholders
            temp_text = text_segment
            placeholders = {}
            placeholder_counter = 0

            for segment in reversed(inline_code_segments):
                placeholder = f"__INLINECODE_{placeholder_counter}__"
                placeholders[placeholder] = segment
                temp_text = temp_text[:segment['start']] + placeholder + temp_text[segment['end']:]
                placeholder_counter += 1

            # Parse from innermost to outermost to handle nesting properly

            # Step 1: Find all bold+italic markers (***text***)
            bold_italic_pattern = r'\*\*\*(.*?)\*\*\*'
            bold_italic_segments = []

            for match in re.finditer(bold_italic_pattern, temp_text):
                bold_italic_segments.append({
                    'start': match.start(),
                    'end': match.end(),
                    'content': match.group(1),
                    'type': 'bold_italic',
                    'original': match.group(0)
                })

            # Step 2: Replace bold+italic with placeholders to avoid conflicts
            for segment in reversed(bold_italic_segments):  # Process from end to avoid position shifts
                placeholder = f"__BOLDITALIC_{placeholder_counter}__"
                placeholders[placeholder] = segment
                temp_text = temp_text[:segment['start']] + placeholder + temp_text[segment['end']:]
                placeholder_counter += 1

            # Step 3: Find bold patterns (avoiding placeholders)
            bold_pattern = r'\*\*((?:(?!\*\*).)*?)\*\*'
            bold_segments = []

            for match in re.finditer(bold_pattern, temp_text):
                # Skip if this is inside a placeholder
                match_text = match.group(0)
                if not any(placeholder in match_text for placeholder in placeholders.keys()):
                    bold_segments.append({
                        'start': match.start(),
                        'end': match.end(),
                        'content': match.group(1),
                        'type': 'bold',
                        'original': match.group(0)
                    })

            # Step 4: Replace bold with placeholders
            for segment in reversed(bold_segments):
                placeholder = f"__BOLD_{placeholder_counter}__"
                placeholders[placeholder] = segment
                temp_text = temp_text[:segment['start']] + placeholder + temp_text[segment['end']:]
                placeholder_counter += 1

            # Step 5: Find italic patterns (avoiding placeholders and double asterisks)
            italic_pattern = r'(?<!\*)\*([^*]+?)\*(?!\*)'
            italic_segments = []

            for match in re.finditer(italic_pattern, temp_text):
                # Skip if this is inside a placeholder
                match_text = match.group(0)
                if not any(placeholder in match_text for placeholder in placeholders.keys()):
                    italic_segments.append({
                        'start': match.start(),
                        'end': match.end(),
                        'content': match.group(1),
                        'type': 'italic',
                        'original': match.group(0)
                    })

            # Step 6: Replace italic with placeholders
            for segment in reversed(italic_segments):
                placeholder = f"__ITALIC_{placeholder_counter}__"
                placeholders[placeholder] = segment
                temp_text = temp_text[:segment['start']] + placeholder + temp_text[segment['end']:]
                placeholder_counter += 1

            # Step 7: Process nested bold patterns inside bold+italic
            for placeholder, segment in placeholders.items():
                if segment['type'] == 'bold_italic':
                    # Check if there are nested formatting inside bold+italic content
                    nested_content = segment['content']

                    # Find any bold (**) patterns inside
                    nested_bold_pattern = r'\*\*(.*?)\*\*'
                    nested_bold_matches = list(re.finditer(nested_bold_pattern, nested_content))

                    # Find any italic (*) patterns inside
                    nested_italic_pattern = r'(?<!\*)\*([^*]+?)\*(?!\*)'
                    nested_italic_matches = list(re.finditer(nested_italic_pattern, nested_content))

                    if nested_bold_matches or nested_italic_matches:
                        # There's nested formatting - process it
                        processed_content = nested_content

                        # Process nested bold first (from end to avoid position shifts)
                        for match in reversed(nested_bold_matches):
                            nested_placeholder = f"__NESTED_BOLD_{placeholder_counter}__"
                            placeholders[nested_placeholder] = {
                                'content': match.group(1),
                                'type': 'nested_bold_in_bolditalic',
                                'parent': placeholder
                            }
                            processed_content = processed_content[:match.start()] + nested_placeholder + processed_content[match.end():]
                            placeholder_counter += 1

                        # Process nested italic
                        for match in reversed(nested_italic_matches):
                            nested_placeholder = f"__NESTED_ITALIC_{placeholder_counter}__"
                            placeholders[nested_placeholder] = {
                                'content': match.group(1),
                                'type': 'nested_italic_in_bolditalic',
                                'parent': placeholder
                            }
                            processed_content = processed_content[:match.start()] + nested_placeholder + processed_content[match.end():]
                            placeholder_counter += 1

                        # Update the segment content
                        segment['content'] = processed_content

            return temp_text, placeholders

        # Main processing starts here
        processed_text, placeholders = parse_nested_formatting(text)

        # Step 8: Process the text sequentially, replacing placeholders with formatted runs
        current_pos = 0

        # Split processed text by placeholders and normal text
        parts = []

        # Find all placeholder positions
        placeholder_positions = []
        for placeholder in placeholders.keys():
            pos = processed_text.find(placeholder)
            if pos != -1:
                placeholder_positions.append({
                    'start': pos,
                    'end': pos + len(placeholder),
                    'placeholder': placeholder
                })

        # Sort by position
        placeholder_positions.sort(key=lambda x: x['start'])

        # Build parts list
        for pos_info in placeholder_positions:
            # Add normal text before placeholder
            if current_pos < pos_info['start']:
                normal_text = processed_text[current_pos:pos_info['start']]
                if normal_text:
                    parts.append({
                        'type': 'normal',
                        'content': normal_text
                    })

            # Add placeholder
            parts.append({
                'type': 'placeholder',
                'placeholder': pos_info['placeholder'],
                'segment': placeholders[pos_info['placeholder']]
            })

            current_pos = pos_info['end']

        # Add remaining normal text
        if current_pos < len(processed_text):
            remaining_text = processed_text[current_pos:]
            if remaining_text:
                parts.append({
                    'type': 'normal',
                    'content': remaining_text
                })

        # Step 9: Render all parts to the paragraph
        for part in parts:
            if part['type'] == 'normal':
                add_text_run(part['content'])

            elif part['type'] == 'placeholder':
                segment = part['segment']

                if segment['type'] == 'inline_code':
                    # Inline code
                    add_text_run(segment['content'], code=True)

                elif segment['type'] == 'bold_italic':
                    # Check if this has nested content
                    content = segment['content']
                    if '__NESTED_' in content:
                        # Process nested content
                        nested_current_pos = 0

                        # Find nested placeholders in content
                        nested_placeholders = [p for p in placeholders.keys() if p in content and 'NESTED_' in p]

                        for nested_placeholder in nested_placeholders:
                            nested_pos = content.find(nested_placeholder)
                            if nested_pos != -1:
                                # Add text before nested
                                if nested_current_pos < nested_pos:
                                    before_text = content[nested_current_pos:nested_pos]
                                    if before_text:
                                        add_text_run(before_text, bold=True, italic=True)

                                # Add nested content
                                nested_segment = placeholders[nested_placeholder]
                                if nested_segment['type'] == 'nested_bold_in_bolditalic':
                                    add_text_run(nested_segment['content'], bold=True, italic=True)
                                elif nested_segment['type'] == 'nested_italic_in_bolditalic':
                                    add_text_run(nested_segment['content'], bold=True, italic=True)

                                nested_current_pos = nested_pos + len(nested_placeholder)

                        # Add remaining text
                        if nested_current_pos < len(content):
                            remaining = content[nested_current_pos:]
                            if remaining:
                                add_text_run(remaining, bold=True, italic=True)
                    else:
                        # Simple bold+italic
                        add_text_run(content, bold=True, italic=True)

                elif segment['type'] == 'bold':
                    # Check for nested italic inside bold
                    content = segment['content']
                    if '*' in content and not content.startswith('*'):  # Has italic inside
                        # Process italic inside bold
                        italic_inside_pattern = r'(?<!\*)\*([^*]+?)\*(?!\*)'
                        last_pos = 0

                        for italic_match in re.finditer(italic_inside_pattern, content):
                            # Add bold text before italic
                            if last_pos < italic_match.start():
                                add_text_run(content[last_pos:italic_match.start()], bold=True)

                            # Add bold+italic text
                            add_text_run(italic_match.group(1), bold=True, italic=True)
                            last_pos = italic_match.end()

                        # Add remaining bold text
                        if last_pos < len(content):
                            add_text_run(content[last_pos:], bold=True)
                    else:
                        # Simple bold
                        add_text_run(content, bold=True)

                elif segment['type'] == 'italic':
                    # Check for nested bold inside italic
                    content = segment['content']
                    if '**' in content:  # Has bold inside
                        # Process bold inside italic
                        bold_inside_pattern = r'\*\*(.*?)\*\*'
                        last_pos = 0

                        for bold_match in re.finditer(bold_inside_pattern, content):
                            # Add italic text before bold
                            if last_pos < bold_match.start():
                                add_text_run(content[last_pos:bold_match.start()], italic=True)

                            # Add bold+italic text
                            add_text_run(bold_match.group(1), bold=True, italic=True)
                            last_pos = bold_match.end()

                        # Add remaining italic text
                        if last_pos < len(content):
                            add_text_run(content[last_pos:], italic=True)
                    else:
                        # Simple italic
                        add_text_run(content, italic=True)

    def flush_code_block(self):
        """Flush accumulated code block lines to document"""
        if self.code_block_lines:
            # Create code block paragraph with all lines
            code_content = '\n'.join(self.code_block_lines)
            p = self.document.add_paragraph(code_content)
            p.style = self.styles_manager.get_style('code_block')

            # Reset code block state
            self.code_block_lines = []
            self.code_language = ""
            self.in_code_block = False

    def parse_line(self, line):
        """Parse a single markdown line and add to document"""
        # Handle code blocks first
        if line.strip().startswith('```'):
            if not self.in_code_block:
                # Start of code block
                self.in_code_block = True
                # Extract language if present (e.g., ```bash)
                lang_match = re.match(r'^```(\w+)?', line.strip())
                self.code_language = lang_match.group(1) if lang_match and lang_match.group(1) else ""
                return
            else:
                # End of code block
                self.flush_code_block()
                return

        # If we're inside a code block, accumulate lines
        if self.in_code_block:
            self.code_block_lines.append(line)
            return

        # Normal markdown processing
        line = line.strip()

        if not line:
            p = self.document.add_paragraph()
            p.style = self.styles_manager.get_style('normal')
            return

        # Headers - preserve formatting in content
        if line.startswith('# '):
            content = line[2:]
            p = self.document.add_paragraph()
            p.style = self.styles_manager.get_style('h1')
            self.process_formatted_text(content, p)

        elif line.startswith('## '):
            content = line[3:]
            p = self.document.add_paragraph()
            p.style = self.styles_manager.get_style('h2')
            self.process_formatted_text(content, p)

        elif line.startswith('### '):
            content = line[4:]
            p = self.document.add_paragraph()
            p.style = self.styles_manager.get_style('h3')
            self.process_formatted_text(content, p)

        elif line.startswith('#### '):
            content = line[5:]
            p = self.document.add_paragraph()
            p.style = self.styles_manager.get_style('h4')
            self.process_formatted_text(content, p)

        elif line.startswith('##### '):
            content = line[6:]
            p = self.document.add_paragraph()
            p.style = self.styles_manager.get_style('h5')
            self.process_formatted_text(content, p)

        elif line.startswith('###### '):
            content = line[7:]
            p = self.document.add_paragraph()
            p.style = self.styles_manager.get_style('h6')
            self.process_formatted_text(content, p)

        # Blockquotes
        elif line.startswith('> '):
            content = line[2:]
            p = self.document.add_paragraph()
            p.style = self.styles_manager.get_style('quote')
            self.process_formatted_text(content, p)

        # Horizontal rule
        elif line.startswith('---') or line.startswith('***'):
            p = self.document.add_paragraph('_' * 50)
            p.style = self.styles_manager.get_style('horizontal_rule')  # Use unified horizontal rule style

        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            content = line[2:]
            p = self.document.add_paragraph()
            p.style = self.styles_manager.get_style('list')  # Use unified list style
            p.add_run('• ')  # Add bullet
            self.process_formatted_text(content, p)

        # Ordered list
        elif re.match(r'^\d+\.\s', line):
            number = re.match(r'^(\d+)\.', line).group(1)
            content = re.sub(r'^\d+\.\s', '', line)
            p = self.document.add_paragraph()
            p.style = self.styles_manager.get_style('list')  # Use unified list style
            p.add_run(f"{number}. ")
            self.process_formatted_text(content, p)

        # Profile section
        elif line.startswith('**PROFILE**') or line.startswith('### **PROFILE**'):
            p = self.document.add_paragraph("PROFILE")
            p.style = self.styles_manager.get_style('h3')

        # Date/episode markers
        elif re.match(r'^\*\([0-9.]+\)\*$', line):
            content = line.replace('*', '')
            p = self.document.add_paragraph(content)
            p.style = self.styles_manager.get_style('profile')

        # Regular paragraph
        else:
            if line.strip():
                p = self.document.add_paragraph()
                if any(keyword in line.lower() for keyword in ['born', 'representative works', 'website']):
                    p.style = self.styles_manager.get_style('profile')
                else:
                    p.style = self.styles_manager.get_style('normal')
                self.process_formatted_text(line, p)


class MarkdownToDocxProcessor:
    """Main processor for converting markdown articles to single DOCX"""

    def __init__(self):
        self.articles = []
        self.article_page_counts = {}
        self.styles_manager = None

    def scan_markdown_files(self, directory_path):
        """Scan directory for markdown files and extract article info"""
        directory = Path(directory_path)

        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")

        # Look for article_XXX_eng.md pattern and sort numerically
        md_files = []
        for md_file in directory.glob("article_*_eng.md"):
            match = re.search(r'article_(\d+)_eng\.md', md_file.name)
            if match:
                article_num = int(match.group(1))
                md_files.append((article_num, md_file))

        md_files.sort(key=lambda x: x[0])

        articles = []
        for article_num, md_file in md_files:
            try:
                with open(md_file, 'r', encoding='utf-8') as f:
                    content = f.read()

                title = f"Episode {article_num}"
                lines = content.split('\n')

                for line in lines[:20]:
                    stripped_line = line.strip()
                    if stripped_line.startswith('## '):
                        potential_title = stripped_line[3:].strip()
                        if len(potential_title) > 5:
                            title = potential_title
                            title = re.sub(r'\*+', '', title)
                            title = title.replace('**', '').replace('*', '')
                            break

                articles.append({
                    'file_path': md_file,
                    'article_number': article_num,
                    'title': title,
                    'content': content
                })

            except Exception as e:
                print(f"WARNING: Error reading {md_file}: {e}")

        self.articles = articles
        print(f"INFO: Successfully loaded {len(self.articles)} articles")
        return self.articles

    def generate_single_article_docx(self, article, output_dir):
        """Generate individual DOCX for a single article and measure its page count"""
        try:
            # Create individual document with styles
            doc = Document()
            styles_manager = create_styles_for_document(doc, verbose=False)
            parser = MarkdownDocxParser(doc, styles_manager)

            # Parse article content
            lines = article['content'].split('\n')
            for line in lines:
                parser.parse_line(line)

            # Flush any remaining code block
            parser.flush_code_block()

            # Save individual DOCX file
            single_docx = output_dir / f"article_{article['article_number']:03d}.docx"
            doc.save(str(single_docx))

            # Convert DOCX to PDF using LibreOffice
            pdf_path = self.convert_docx_to_pdf_libreoffice(single_docx, output_dir)

            if pdf_path:
                # Count pages in the generated PDF
                page_count = self.count_pdf_pages(pdf_path)
                print(f"INFO: Article {article['article_number']:03d}: {page_count} pages (PDF: {Path(pdf_path).name})")
                return page_count
            else:
                print(f"ERROR: Could not convert article {article['article_number']:03d} to PDF")
                return 1  # Default to 1 page on error

        except Exception as e:
            print(f"ERROR: Failed to generate single article {article['article_number']}: {e}")
            return 1

    def convert_docx_to_pdf_libreoffice(self, docx_path, output_dir):
        """Convert DOCX to PDF using LibreOffice headless (exact conversion)"""
        try:
            docx_path = Path(docx_path)
            output_dir = Path(output_dir)

            # Calculate PDF output path
            pdf_path = output_dir / f"{docx_path.stem}.pdf"

            # Simple LibreOffice command
            cmd = [
                'libreoffice',
                '--headless',
                '--convert-to', 'pdf',
                str(docx_path.name)  # Just the filename, not full path
            ]

            # Run conversion from the directory where files are located
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=90,
                cwd=str(output_dir)  # Work from the single_articles directory
            )

            # Check if PDF was created
            if pdf_path.exists():
                return str(pdf_path)
            else:
                print(f"ERROR: PDF not created for {docx_path.name}")
                return None

        except subprocess.TimeoutExpired:
            print(f"ERROR: LibreOffice timeout for {docx_path}")
            return None
        except Exception as e:
            print(f"ERROR: LibreOffice failed: {e}")
            return None

    def count_pdf_pages(self, pdf_path):
        """Count pages in a PDF file using pypdf"""
        try:
            reader = PdfReader(pdf_path)
            page_count = len(reader.pages)
            return page_count
        except Exception as e:
            print(f"ERROR: Could not count pages in {pdf_path}: {e}")
            return 1  # Default to 1 page on error

    def generate_all_single_articles(self, output_dir):
        """Generate individual DOCX files for each article and measure page counts"""
        print("INFO: Generating individual article DOCX files...")

        # Create output directory
        single_articles_dir = output_dir / "00_single_articles"
        single_articles_dir.mkdir(parents=True, exist_ok=True)

        page_counts = {}

        for i, article in enumerate(self.articles):
            print(f"INFO: Processing single article {i+1}/{len(self.articles)}: {article['title'][:50]}...")

            page_count = self.generate_single_article_docx(article, single_articles_dir)
            page_counts[i] = page_count

        self.article_page_counts = page_counts
        print(f"INFO: All single articles generated in {single_articles_dir}")

        # Save page counts to JSON
        self.save_page_counts_json(output_dir)

        return page_counts

    def calculate_exact_page_numbers(self, index_offset, title_offset):
        """Calculate exact page numbers based on actual article page counts"""
        page_mapping = {}
        current_page = title_offset + index_offset + 1

        print(f"INFO: Calculating exact pages - First article starts at page {current_page}")

        for i, article in enumerate(self.articles):
            page_mapping[i] = current_page

            # Use actual measured page count
            article_pages = self.article_page_counts.get(i, 1)
            current_page += article_pages

            article_num = article.get('article_number', i+1)
            print(f"INFO: Article {article_num:03d} -> Page {page_mapping[i]} ({article_pages} pages)")

        return page_mapping

    def save_page_counts_json(self, output_dir):
        """Save page counts and article info to JSON file"""
        try:
            page_data = []

            for i, article in enumerate(self.articles):
                page_count = self.article_page_counts.get(i, 1)
                article_num = article.get('article_number', i+1)

                page_data.append({
                    'article_number': article_num,
                    'title': article['title'],
                    'pages': page_count,
                    'file_path': str(article['file_path'])
                })

            # Save to JSON
            json_path = output_dir / "articles_page_counts.json"
            import json

            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(page_data, f, indent=2, ensure_ascii=False)

            print(f"INFO: Page counts saved to: {json_path}")

        except Exception as e:
            print(f"ERROR: Failed to save page counts JSON: {e}")

    def create_cover_page(self, document, styles_manager):
        """Create cover page using external styles"""
        p = document.add_paragraph(styles_manager.config.document_info['title'])
        p.style = styles_manager.get_style('cover_title')

        p = document.add_paragraph(styles_manager.config.document_info['subtitle'])
        p.style = styles_manager.get_style('cover_subtitle')

        # Author with Japanese
        p = document.add_paragraph()
        p.style = styles_manager.get_style('cover_subtitle')
        run1 = p.add_run(f"{styles_manager.config.document_info['author']} (")
        run2 = p.add_run(styles_manager.config.document_info['author_japanese'])
        run2.font.name = styles_manager.document_styles.font_manager.get_font('japanese')
        run3 = p.add_run(")")

        p = document.add_paragraph(f"{len(self.articles)} Episodes")
        p.style = styles_manager.get_style('cover_subtitle')

        p = document.add_paragraph(f"Generated: {styles_manager.config.get_generation_date()}")
        p.style = styles_manager.get_style('cover_subtitle')

        document.add_page_break()

    def create_blank_page(self, document, styles_manager):
        """Create intentionally blank page using config spacing"""
        for _ in range(15):
            document.add_paragraph()

        p = document.add_paragraph("This page intentionally left blank")
        p.style = styles_manager.get_style('blank')  # Use unified blank style

        document.add_page_break()

    def create_table_of_contents(self, document, styles_manager, page_mapping):
        """Create table of contents with exact page numbers"""
        p = document.add_paragraph("TABLE OF CONTENTS")
        p.style = styles_manager.get_style('h1')

        document.add_paragraph()  # Spacing

        for i, article in enumerate(self.articles):
            title = article['title'].replace('**', '').replace('*', '').strip()
            article_num = article.get('article_number', i+1)
            article_page = page_mapping.get(i, 1)

            # Full titles without truncation
            page_str = str(article_page)

            # Intelligent dot spacing
            base_line_length = 75
            dots_needed = max(5, base_line_length - len(title) - len(page_str))

            if len(title) > 60:
                dots_needed = max(3, 15)

            dots = '.' * min(dots_needed, 40)
            toc_text = f"{title} {dots} {page_str}"

            p = document.add_paragraph(toc_text)
            p.style = styles_manager.get_style('toc')

            if i < 5:
                print(f"DEBUG: TOC {article_num:03d}: Page {article_page} - {title[:50]}...")

        document.add_page_break()

    def merge_articles_into_final_docx(self, output_path, index_offset, title_offset):
        """Create final DOCX by merging individual articles with exact page numbers"""
        print("INFO: Creating final merged document...")

        # Create final document with styles
        final_doc = Document()
        styles_manager = create_styles_for_document(final_doc, verbose=True)

        # Calculate exact page numbers based on measured page counts
        page_mapping = self.calculate_exact_page_numbers(index_offset, title_offset)

        # 1. Cover page
        print("INFO: Adding cover page...")
        self.create_cover_page(final_doc, styles_manager)

        # 2. Blank page
        print("INFO: Adding blank page...")
        self.create_blank_page(final_doc, styles_manager)

        # 3. Table of contents with exact pages
        print("INFO: Adding table of contents with exact page numbers...")
        self.create_table_of_contents(final_doc, styles_manager, page_mapping)

        # 4. Add all articles
        print("INFO: Adding all articles...")
        parser = MarkdownDocxParser(final_doc, styles_manager)

        for i, article in enumerate(self.articles):
            print(f"INFO: Adding article {i+1}/{len(self.articles)}: {article['title'][:50]}...")

            # Parse and add article content
            lines = article['content'].split('\n')
            for line in lines:
                parser.parse_line(line)

            # Flush any remaining code block
            parser.flush_code_block()

            # Add page break after each article (except last)
            if i < len(self.articles) - 1:
                final_doc.add_page_break()

        # Save final document
        final_doc.save(str(output_path))
        print(f"INFO: Final document saved: {output_path}")

        # Cleanup temporary files
        self.cleanup_temporary_files(styles_manager.config.processing['output_dir'])

    def cleanup_temporary_files(self, output_dir):
        """Clean up temporary DOCX and PDF files"""
        single_articles_dir = output_dir / "00_single_articles"

        if single_articles_dir.exists():
            print("INFO: Cleaning up temporary files...")

            # Count files before cleanup
            docx_files = list(single_articles_dir.glob("*.docx"))
            pdf_files = list(single_articles_dir.glob("*.pdf"))

            print(f"INFO: Found {len(docx_files)} DOCX and {len(pdf_files)} PDF files to clean")

            print("INFO: Cleanup completed (files preserved for debugging)")
            print("INFO: To enable cleanup, modify cleanup_temporary_files()")
        else:
            print("INFO: No temporary files to clean")

    def generate_docx(self, markdown_directory, output_path="articles_collection.docx",
                     extract_titles_only=False, index_offset=13, title_offset=2):
        """Generate complete DOCX from markdown articles with exact page calculation"""
        try:
            # Create temporary document to get config
            temp_doc = Document()
            temp_styles = create_styles_for_document(temp_doc)
            output_path = temp_styles.config.processing['output_dir'] / Path(output_path).name
            temp_styles.config.processing['output_dir'].mkdir(exist_ok=True)

            print(f"INFO: Starting DOCX generation: {output_path}")
            print(f"INFO: Index offset: {index_offset} pages, Title offset: {title_offset} pages")

            # Scan markdown files
            self.scan_markdown_files(markdown_directory)

            if extract_titles_only:
                print("INFO: Title extraction completed.")
                return None

            # Generate individual articles only
            print("INFO: Generating individual DOCX files...")
            self.generate_all_single_articles(temp_styles.config.processing['output_dir'])

            print(f"INFO: Individual DOCX generation completed")
            print(f"INFO: Total articles processed: {len(self.articles)}")
            print(f"INFO: Files saved in: ./output/00_single_articles/")

            return "Individual DOCX files generated successfully"

        except Exception as e:
            print(f"ERROR: DOCX generation failed: {e}")
            raise


def main():
    """Main function with command line argument support"""
    parser = argparse.ArgumentParser(
        description='Markdown to DOCX Processor\n\n'
                    'Converts multiple markdown articles to a single DOCX document.\n'
                    'Phase 1: Creates individual DOCX files to measure exact page counts\n'
                    'Phase 2: Creates final merged document with accurate TOC\n\n'
                    'Supported markdown elements:\n'
                    '  • Headers: # ## ### #### ##### ######\n'
                    '  • Blockquotes: > quoted text\n'
                    '  • Bold: **text** and italic: *text*\n'
                    '  • Bold+Italic: ***text*** or **bold *italic* nested**\n'
                    '  • Inline code: `code`\n'
                    '  • Code blocks: ```language\\ncode\\n```\n'
                    '  • Lists: - item or 1. item\n'
                    '  • Japanese character support\n\n'
                    'Custom fonts used:\n'
                    '  • BPG Serif GPL&GNU (titles)\n'
                    '  • BPG Gorda GPL&GNU (headers)\n'
                    '  • Amiri (body text)\n'
                    '  • Noto Sans CJK JP (Japanese characters)\n'
                    '  • Courier (code blocks - uses ./fonts/courier.ttf)\n\n'
                    'Example: python ArticlesDocxNewer.py ./articles/ collection.docx -io 11',
        formatter_class=argparse.RawDescriptionHelpFormatter
    )

    parser.add_argument(
        'markdown_directory',
        help='Directory containing markdown articles (article_001_eng.md format)'
    )

    parser.add_argument(
        'output_docx',
        nargs='?',
        default='articles_collection.docx',
        help='Output DOCX filename (default: articles_collection.docx)'
    )

    parser.add_argument(
        '-io', '--index-offset',
        type=int,
        default=13,
        help='Number of pages the table of contents will occupy (default: 13)'
    )

    parser.add_argument(
        '-to', '--title-offset',
        type=int,
        default=2,
        help='Number of initial pages before TOC (cover + blank, default: 2)'
    )

    parser.add_argument(
        '--extract-titles',
        action='store_true',
        help='Only extract and validate titles, don\'t generate DOCX'
    )

    parser.add_argument(
        '-v', '--verbose',
        action='store_true',
        help='Enable verbose output'
    )

    args = parser.parse_args()

    try:
        if not Path(args.markdown_directory).exists():
            print(f"ERROR: Directory not found: {args.markdown_directory}")
            return 1

        processor = MarkdownToDocxProcessor()

        if args.verbose:
            print(f"Configuration:")
            print(f"  Markdown directory: {args.markdown_directory}")
            print(f"  DOCX output: {args.output_docx}")
            print(f"  Index offset: {args.index_offset} pages")
            print(f"  Title offset: {args.title_offset} pages")
            print()

        output_path = processor.generate_docx(
            markdown_directory=args.markdown_directory,
            output_path=args.output_docx,
            extract_titles_only=args.extract_titles,
            index_offset=args.index_offset,
            title_offset=args.title_offset
        )

        if output_path:
            print(f"SUCCESS: DOCX collection generated at {output_path}")

        return 0

    except KeyboardInterrupt:
        print("\nOperation cancelled by user")
        return 130

    except Exception as e:
        print(f"ERROR: {e}")
        if args.verbose:
            import traceback
            traceback.print_exc()
        return 1


if __name__ == "__main__":
    sys.exit(main())
