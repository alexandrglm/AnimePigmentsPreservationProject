#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Finals.py
Final document assembler - Creates complete DOCX with accurate page numbering
"""

import json
import sys
import argparse
import subprocess
import zipfile
import tempfile
import shutil
from pathlib import Path
from datetime import datetime
from xml.etree import ElementTree as ET

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml

# Import external styles module
from DocxStyles import create_styles_for_document

try:
    from pypdf import PdfReader
except ImportError:
    print("ERROR: pypdf not installed. Install with: pip install pypdf")
    sys.exit(1)


class DocumentMerger:
    """Advanced DOCX merger that preserves all formatting, page breaks, and structure via XML manipulation"""

    def __init__(self):
        self.namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        }
        # Register namespaces for XML parsing
        for prefix, uri in self.namespaces.items():
            ET.register_namespace(prefix, uri)

    def extract_document_xml(self, docx_path):
        """Extract the main document.xml from a DOCX file"""
        try:
            with zipfile.ZipFile(docx_path, 'r') as zip_file:
                # Read document.xml
                document_xml = zip_file.read('word/document.xml')
                return ET.fromstring(document_xml)
        except Exception as e:
            print(f"ERROR: Could not extract XML from {docx_path}: {e}")
            return None

    def get_body_content(self, doc_root):
        """Extract body content from document XML"""
        try:
            # Find w:body element
            body = doc_root.find('.//w:body', self.namespaces)
            if body is not None:
                return list(body)  # Return all child elements
            return []
        except Exception as e:
            print(f"ERROR: Could not extract body content: {e}")
            return []

    def create_page_break_element(self):
        """Create a page break element in Word XML format"""
        # Create: <w:p><w:r><w:br w:type="page"/></w:r></w:p>
        p_elem = ET.Element(f"{{{self.namespaces['w']}}}p")
        r_elem = ET.SubElement(p_elem, f"{{{self.namespaces['w']}}}r")
        br_elem = ET.SubElement(r_elem, f"{{{self.namespaces['w']}}}br")
        br_elem.set(f"{{{self.namespaces['w']}}}type", "page")
        return p_elem

    def ends_with_page_break(self, source_content):
        """Check if the last element of content ends with a page break"""
        try:
            if not source_content:
                return False

            # Check the last few elements for page breaks
            # Sometimes page breaks can be in the last or second-to-last element
            elements_to_check = source_content[-3:] if len(source_content) >= 3 else source_content

            for element in reversed(elements_to_check):
                # Skip sectPr elements
                if element.tag.endswith('}sectPr'):
                    continue

                # Check if this element contains a page break
                for r_elem in element.findall(f'.//{{{self.namespaces["w"]}}}r'):
                    for br_elem in r_elem.findall(f'.//{{{self.namespaces["w"]}}}br'):
                        br_type = br_elem.get(f'{{{self.namespaces["w"]}}}type')
                        if br_type == "page":
                            print(f"DEBUG: Found trailing page break in element {element.tag}")
                            return True

            return False

        except Exception as e:
            print(f"WARNING: Error checking trailing page break: {e}")
            return False

    def ends_with_page_break(self, source_content):
        """Check if the last element of content ends with a page break"""
        try:
            if not source_content:
                return False

            # Check the last few elements for page breaks
            # Sometimes page breaks can be in the last or second-to-last element
            elements_to_check = source_content[-3:] if len(source_content) >= 3 else source_content

            for element in reversed(elements_to_check):
                # Skip sectPr elements
                if element.tag.endswith('}sectPr'):
                    continue

                # Check if this element contains a page break
                for r_elem in element.findall(f'.//{{{self.namespaces["w"]}}}r'):
                    for br_elem in r_elem.findall(f'.//{{{self.namespaces["w"]}}}br'):
                        br_type = br_elem.get(f'{{{self.namespaces["w"]}}}type')
                        if br_type == "page":
                            print(f"DEBUG: Found trailing page break in element {element.tag}")
                            return True

            return False

        except Exception as e:
            print(f"WARNING: Error checking trailing page break: {e}")
            return False

    def has_content_after_last_page_break(self, source_content):
        """Check if there's meaningful content after the last page break"""
        try:
            if not source_content:
                return False

            # Find the last page break position
            last_page_break_index = -1

            for i, element in enumerate(source_content):
                # Skip sectPr elements
                if element.tag.endswith('}sectPr'):
                    continue

                # Check if this element contains a page break
                for r_elem in element.findall(f'.//{{{self.namespaces["w"]}}}r'):
                    for br_elem in r_elem.findall(f'.//{{{self.namespaces["w"]}}}br'):
                        br_type = br_elem.get(f'{{{self.namespaces["w"]}}}type')
                        if br_type == "page":
                            last_page_break_index = i

            if last_page_break_index == -1:
                # No page breaks found, content exists
                return True

            # Check if there's meaningful content after the last page break
            for i in range(last_page_break_index + 1, len(source_content)):
                element = source_content[i]

                # Skip sectPr elements
                if element.tag.endswith('}sectPr'):
                    continue

                # Check for text content
                text_elements = element.findall(f'.//{{{self.namespaces["w"]}}}t')
                has_text = any(elem.text and elem.text.strip() for elem in text_elements)

                if has_text:
                    print(f"DEBUG: Found content after last page break at position {i}")
                    return True

            print(f"DEBUG: No meaningful content after last page break")
            return False

        except Exception as e:
            print(f"WARNING: Error checking content after page break: {e}")
            return True  # Default to safe side
        """Remove ONLY problematic trailing elements that cause blank pages - ultra conservative"""
        try:
            if not source_content:
                return source_content

            # Work on a copy to avoid modifying original
            cleaned_content = source_content.copy()
            removed_count = 0

            # Check only the last 3 elements (ultra conservative)
            elements_to_check = min(3, len(cleaned_content))

            # Remove trailing elements from the end, one by one
            for _ in range(elements_to_check):
                if not cleaned_content:
                    break

                last_element = cleaned_content[-1]

                # Skip sectPr elements (always keep them)
                if last_element.tag.endswith('}sectPr'):
                    break

                # Check if element is truly empty and problematic
                is_problematic = True

                # Check for text content
                text_elements = last_element.findall(f'.//{{{self.namespaces["w"]}}}t')
                has_text = any(elem.text and elem.text.strip() for elem in text_elements)

                if has_text:
                    # Has text, preserve it
                    is_problematic = False

                # Check for page breaks (preserve them)
                br_elements = last_element.findall(f'.//{{{self.namespaces["w"]}}}br')
                has_breaks = any(br.get(f'{{{self.namespaces["w"]}}}type') == "page" for br in br_elements)

                if has_breaks:
                    # Has page breaks, preserve it
                    is_problematic = False

                # Check for images, tables, or other content
                # Look for common content elements
                content_tags = ['drawing', 'tbl', 'r']  # drawing, table, run
                has_other_content = any(last_element.findall(f'.//{{{self.namespaces["w"]}}}{tag}') for tag in content_tags)

                if has_other_content:
                    # Has other content, preserve it
                    is_problematic = False

                # Only remove if truly empty and problematic
                if is_problematic:
                    # Check what type of empty element this is
                    element_info = f"tag: {last_element.tag.split('}')[-1] if '}' in last_element.tag else last_element.tag}"
                    print(f"INFO: Removing problematic trailing element - {element_info}")
                    cleaned_content.pop()  # Remove last element
                    removed_count += 1
                else:
                    # Found meaningful content, stop cleaning
                    break

            if removed_count > 0:
                print(f"INFO: Pre-cleaning removed {removed_count} problematic trailing elements")

            return cleaned_content

        except Exception as e:
            print(f"WARNING: Pre-cleaning failed, using original content: {e}")
            return source_content
        """Check if there's meaningful content after the last page break"""
        try:
            if not source_content:
                return False

            # Find the last page break position
            last_page_break_index = -1

            for i, element in enumerate(source_content):
                # Skip sectPr elements
                if element.tag.endswith('}sectPr'):
                    continue

                # Check if this element contains a page break
                for r_elem in element.findall(f'.//{{{self.namespaces["w"]}}}r'):
                    for br_elem in r_elem.findall(f'.//{{{self.namespaces["w"]}}}br'):
                        br_type = br_elem.get(f'{{{self.namespaces["w"]}}}type')
                        if br_type == "page":
                            last_page_break_index = i

            if last_page_break_index == -1:
                # No page breaks found, content exists
                return True

            # Check if there's meaningful content after the last page break
            for i in range(last_page_break_index + 1, len(source_content)):
                element = source_content[i]

                # Skip sectPr elements
                if element.tag.endswith('}sectPr'):
                    continue

                # Check for text content
                text_elements = element.findall(f'.//{{{self.namespaces["w"]}}}t')
                has_text = any(elem.text and elem.text.strip() for elem in text_elements)

                if has_text:
                    print(f"DEBUG: Found content after last page break at position {i}")
                    return True

            print(f"DEBUG: No meaningful content after last page break")
            return False

        except Exception as e:
            print(f"WARNING: Error checking content after page break: {e}")
            return True  # Default to safe side
        """Create a page break element in Word XML format"""
        # Create: <w:p><w:r><w:br w:type="page"/></w:r></w:p>
        p_elem = ET.Element(f"{{{self.namespaces['w']}}}p")
        r_elem = ET.SubElement(p_elem, f"{{{self.namespaces['w']}}}r")
        br_elem = ET.SubElement(r_elem, f"{{{self.namespaces['w']}}}br")
        br_elem.set(f"{{{self.namespaces['w']}}}type", "page")
        return p_elem

    def is_page_break_paragraph(self, p_element):
        """Check if a paragraph element contains only a page break"""
        try:
            # Look for w:p > w:r > w:br[@w:type="page"]
            for r_elem in p_element.findall(f'.//{{{self.namespaces["w"]}}}r'):
                for br_elem in r_elem.findall(f'.//{{{self.namespaces["w"]}}}br'):
                    br_type = br_elem.get(f'{{{self.namespaces["w"]}}}type')
                    if br_type == "page":
                        # Check if this paragraph has ONLY the page break (no text)
                        text_elements = p_element.findall(f'.//{{{self.namespaces["w"]}}}t')
                        has_text = any(elem.text and elem.text.strip() for elem in text_elements)
                        return not has_text  # True if no text, only page break
            return False
        except Exception as e:
            print(f"WARNING: Error checking page break: {e}")
            return False

    def remove_duplicate_page_breaks(self, body_element):
        """Remove consecutive page breaks to prevent blank pages"""
        try:
            print("INFO: Post-processing XML to remove duplicate page breaks...")

            # Get all paragraph elements in body
            paragraphs = body_element.findall(f'{{{self.namespaces["w"]}}}p')

            removed_count = 0
            i = 0

            while i < len(paragraphs) - 1:
                current_p = paragraphs[i]
                next_p = paragraphs[i + 1]

                # Check if both are page break paragraphs
                if (self.is_page_break_paragraph(current_p) and
                    self.is_page_break_paragraph(next_p)):

                    print(f"INFO: Found duplicate page breaks at position {i}, removing second one")

                    # Remove the second page break from body
                    body_element.remove(next_p)

                    # Update paragraphs list (re-fetch after removal)
                    paragraphs = body_element.findall(f'{{{self.namespaces["w"]}}}p')
                    removed_count += 1

                    # Don't increment i, check same position again
                    continue

                i += 1

            if removed_count > 0:
                print(f"INFO: Removed {removed_count} duplicate page breaks")
            else:
                print("INFO: No duplicate page breaks found")

            return removed_count

        except Exception as e:
            print(f"ERROR: Failed to remove duplicate page breaks: {e}")
            return 0

    def remove_trailing_empty_paragraphs(self, body_element):
        """Remove empty paragraphs that might cause spacing issues"""
        try:
            paragraphs = body_element.findall(f'{{{self.namespaces["w"]}}}p')
            removed_count = 0

            # Remove empty paragraphs from the end
            while paragraphs:
                last_p = paragraphs[-1]

                # Check if paragraph is empty (no text, no page breaks, no special content)
                text_elements = last_p.findall(f'.//{{{self.namespaces["w"]}}}t')
                has_text = any(elem.text and elem.text.strip() for elem in text_elements)

                br_elements = last_p.findall(f'.//{{{self.namespaces["w"]}}}br')
                has_breaks = len(br_elements) > 0

                if not has_text and not has_breaks:
                    # This is an empty paragraph, remove it
                    body_element.remove(last_p)
                    paragraphs = body_element.findall(f'{{{self.namespaces["w"]}}}p')
                    removed_count += 1
                else:
                    # Found non-empty paragraph, stop
                    break

            if removed_count > 0:
                print(f"INFO: Removed {removed_count} trailing empty paragraphs")

            return removed_count

        except Exception as e:
            print(f"ERROR: Failed to remove trailing paragraphs: {e}")
            return 0

    def merge_docx_xml_advanced(self, target_docx_path, source_docx_paths, output_path):
        """
        Advanced XML-level merger that preserves all formatting and page breaks

        Args:
            target_docx_path: Base DOCX file (will be modified)
            source_docx_paths: List of DOCX files to merge
            output_path: Final output path
        """
        try:
            print("INFO: Starting advanced XML merger...")

            # Create temporary working directory
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_dir = Path(temp_dir)

                # Extract target DOCX to temp directory
                target_zip_path = temp_dir / "target.zip"
                shutil.copy2(target_docx_path, target_zip_path)

                # Extract ZIP
                with zipfile.ZipFile(target_zip_path, 'r') as zip_file:
                    zip_file.extractall(temp_dir / "docx_content")

                docx_content_dir = temp_dir / "docx_content"

                # Load target document XML
                target_xml_path = docx_content_dir / "word" / "document.xml"
                target_tree = ET.parse(target_xml_path)
                target_root = target_tree.getroot()
                target_body = target_root.find('.//w:body', self.namespaces)

                if target_body is None:
                    raise ValueError("Could not find body element in target document")

                # Process each source document
                for i, source_path in enumerate(source_docx_paths):
                    print(f"INFO: Merging {Path(source_path).name}...")

                    # Extract source XML
                    source_root = self.extract_document_xml(source_path)
                    if source_root is None:
                        print(f"WARNING: Skipping {source_path} - could not extract XML")
                        continue

                    # Get source body content
                    source_content = self.get_body_content(source_root)

                    # PRE-CLEANING: Remove problematic trailing elements
                    cleaned_content = self.clean_trailing_elements(source_content)
                    if len(cleaned_content) != len(source_content):
                        removed = len(source_content) - len(cleaned_content)
                        print(f"INFO: Pre-cleaned {Path(source_path).name} - removed {removed} problematic elements")

                    # Add cleaned content to target body
                    for element in cleaned_content:
                        # Skip sectPr (section properties) as they can cause conflicts
                        if element.tag.endswith('}sectPr'):
                            continue
                        target_body.append(element)

                        # INTELLIGENT PAGE BREAK LOGIC
                        if i < len(source_docx_paths) - 1:  # Not the last document

                            # Check if this document needs a page break (use cleaned content)
                            needs_page_break = True

                            # Method 1: Check if document ends with page break
                            if self.ends_with_page_break(cleaned_content):
                                print(f"🚨 DESFASE DETECTADO: {Path(source_path).name} - Duplicate page break would create blank page")
                                needs_page_break = False

                            # Method 2: Check if there's content after last page break
                            elif not self.has_content_after_last_page_break(cleaned_content):
                                print(f"🚨 DESFASE DETECTADO: {Path(source_path).name} - Duplicate page break would create blank page")
                                needs_page_break = False

                            # Add page break only if needed
                            if needs_page_break:
                                page_break = self.create_page_break_element()
                                target_body.append(page_break)
                                print(f"INFO: Added necessary page break after {Path(source_path).name}")
                            else:
                                print(f"INFO: Skipped page break for {Path(source_path).name} (would create blank page)")

                # POST-PROCESSING: Additional cleanup for any remaining issues
                print("INFO: Starting post-processing to clean up document...")
                duplicate_breaks_removed = self.remove_duplicate_page_breaks(target_body)
                empty_paragraphs_removed = self.remove_trailing_empty_paragraphs(target_body)

                print(f"INFO: Post-processing complete - Removed {duplicate_breaks_removed} duplicate breaks, {empty_paragraphs_removed} empty paragraphs")

                # Save modified document.xml
                target_tree.write(target_xml_path, encoding='utf-8', xml_declaration=True)

                # Recreate DOCX file
                output_path = Path(output_path)
                with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                    # Add all files from extracted content
                    for file_path in docx_content_dir.rglob('*'):
                        if file_path.is_file():
                            # Calculate archive path (relative to docx_content)
                            archive_path = file_path.relative_to(docx_content_dir)
                            zip_file.write(file_path, archive_path)

                print(f"INFO: Advanced XML merger completed: {output_path}")
                return True

        except Exception as e:
            print(f"ERROR: Advanced XML merger failed: {e}")
            return False

    def merge_docx_documents(self, target_doc, source_docx_path):
        """Legacy merger - kept for compatibility but not used in Finals"""
        try:
            source_doc = Document(source_docx_path)

            for paragraph in source_doc.paragraphs:
                if paragraph.text.strip():  # Only copy non-empty paragraphs
                    self.copy_paragraph_with_style(paragraph, target_doc)

            print(f"INFO: Merged content from {Path(source_docx_path).name}")
            return True

        except Exception as e:
            print(f"ERROR: Failed to merge {source_docx_path}: {e}")
            return False


class Finals:
    """Main class for final document assembly"""

    def __init__(self):
        self.page_counts_data = []
        self.articles_dir = Path("./output/00_single_articles/")
        self.output_dir = Path("./output/")
        self.page_counts_file = self.output_dir / "articles_page_counts.json"
        self.merger = DocumentMerger()

    def load_page_counts_json(self):
        """Load page counts data from JSON file"""
        try:
            if not self.page_counts_file.exists():
                raise FileNotFoundError(f"Page counts file not found: {self.page_counts_file}")

            with open(self.page_counts_file, 'r', encoding='utf-8') as f:
                self.page_counts_data = json.load(f)

            print(f"INFO: Loaded page counts for {len(self.page_counts_data)} articles")

            # Show first few entries for verification
            for i, article in enumerate(self.page_counts_data[:5]):
                print(f"  Article {article['article_number']:03d}: {article['pages']} pages - {article['title'][:40]}...")

            return True

        except Exception as e:
            print(f"ERROR: Failed to load page counts: {e}")
            return False

    def validate_prerequisites(self):
        """Validate that all required files exist"""
        print("INFO: Validating prerequisites...")

        missing_files = []

        # Check page counts JSON
        if not self.page_counts_file.exists():
            missing_files.append(str(self.page_counts_file))

        # Check articles directory
        if not self.articles_dir.exists():
            missing_files.append(str(self.articles_dir))

        # Check individual article files
        missing_articles = []
        for article in self.page_counts_data:
            article_file = self.articles_dir / f"article_{article['article_number']:03d}.docx"
            if not article_file.exists():
                missing_articles.append(str(article_file))

        if missing_files:
            print(f"ERROR: Missing required files: {missing_files}")
            return False

        if missing_articles:
            print(f"ERROR: Missing article files: {missing_articles[:5]}...")  # Show first 5
            print(f"Total missing articles: {len(missing_articles)}")
            return False

        print(f"INFO: All prerequisites validated - {len(self.page_counts_data)} articles ready")
        return True

    def measure_title_blank_pages(self):
        """Create title+blank pages, convert to PDF, and measure actual pages"""
        print("INFO: Measuring title and blank pages...")

        try:
            # Create document with styles
            doc = Document()
            styles_manager = create_styles_for_document(doc, verbose=False)

            # Title page (without generation date)
            p = doc.add_paragraph(styles_manager.config.document_info['title'])
            p.style = styles_manager.get_style('cover_title')

            p = doc.add_paragraph(styles_manager.config.document_info['subtitle'])
            p.style = styles_manager.get_style('cover_subtitle')

            # Author with Japanese
            p = doc.add_paragraph()
            p.style = styles_manager.get_style('cover_subtitle')
            run1 = p.add_run(f"{styles_manager.config.document_info['author']} (")
            run2 = p.add_run(styles_manager.config.document_info['author_japanese'])
            run2.font.name = styles_manager.document_styles.font_manager.get_font('japanese')
            run3 = p.add_run(")")

            p = doc.add_paragraph(f"{len(self.page_counts_data)} Episodes")
            p.style = styles_manager.get_style('cover_subtitle')

            # Page break to blank page
            doc.add_page_break()

            # Blank page
            for _ in range(15):
                doc.add_paragraph()

            p = doc.add_paragraph("This page intentionally left blank")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(styles_manager.config.font_sizes['small'])
            p.runs[0].font.color.rgb = styles_manager.config.colors['light_text']

            # Save and measure
            title_docx_path = self.output_dir / "00_title_blank_measure.docx"
            doc.save(str(title_docx_path))

            # Convert to PDF and count pages
            title_pdf_path = self.convert_docx_to_pdf(title_docx_path)
            if title_pdf_path:
                title_pages = self.count_pdf_pages(title_pdf_path)
                print(f"INFO: Title + Blank pages: {title_pages} pages (measured)")
                return title_pages, title_docx_path
            else:
                print("ERROR: Could not measure title pages, using default")
                return 2, title_docx_path

        except Exception as e:
            print(f"ERROR: Failed to measure title pages: {e}")
            return 2, None

    def measure_toc_pages(self, use_real_page_numbers=False, page_mapping=None):
        """Create TOC, convert to PDF, and measure actual pages"""
        if use_real_page_numbers:
            print("INFO: Measuring final TOC with real page numbers...")
        else:
            print("INFO: Measuring preliminary TOC with placeholder numbers...")

        try:
            # Create document with styles
            doc = Document()
            styles_manager = create_styles_for_document(doc, verbose=False)

            # Add TOC title
            p = doc.add_paragraph("TABLE OF CONTENTS")
            p.style = styles_manager.get_style('h1')
            doc.add_paragraph()  # Spacing

            # Add all entries
            for i, article in enumerate(self.page_counts_data):
                title = article['title'].replace('**', '').replace('*', '').strip()

                # Use real page numbers if provided, otherwise placeholder
                if use_real_page_numbers and page_mapping:
                    page_str = str(page_mapping.get(i, 999))
                else:
                    page_str = "XXX"  # Placeholder

                # Intelligent dot spacing
                base_line_length = 75
                dots_needed = max(5, base_line_length - len(title) - len(page_str))
                if len(title) > 60:
                    dots_needed = max(3, 15)

                dots = '.' * min(dots_needed, 40)
                toc_text = f"{title} {dots} {page_str}"

                p = doc.add_paragraph(toc_text)
                p.style = styles_manager.get_style('toc')

            # Save TOC
            if use_real_page_numbers:
                toc_docx_path = self.output_dir / "00_final_toc_measure.docx"
            else:
                toc_docx_path = self.output_dir / "00_preliminary_toc_measure.docx"

            doc.save(str(toc_docx_path))

            # Convert to PDF and count pages
            toc_pdf_path = self.convert_docx_to_pdf(toc_docx_path)
            if toc_pdf_path:
                toc_pages = self.count_pdf_pages(toc_pdf_path)
                if use_real_page_numbers:
                    print(f"INFO: Final TOC pages: {toc_pages} pages (measured)")
                else:
                    print(f"INFO: Preliminary TOC pages: {toc_pages} pages (measured)")
                return toc_pages, toc_docx_path
            else:
                print("ERROR: Could not measure TOC pages, using default")
                return 13, toc_docx_path

        except Exception as e:
            print(f"ERROR: Failed to measure TOC: {e}")
            return 13, None

    def calculate_exact_pages(self, title_pages, toc_pages):
        """Calculate exact page numbers for each article"""
        print(f"INFO: Calculating exact pages - Title: {title_pages}, TOC: {toc_pages}")

        page_mapping = {}
        current_page = 1

        # Title + blank pages
        current_page += title_pages

        # TOC pages
        current_page += toc_pages

        print(f"INFO: Articles start at page {current_page}")

        # Calculate page for each article
        for i, article in enumerate(self.page_counts_data):
            page_mapping[i] = current_page
            article_pages = article['pages']
            current_page += article_pages

            print(f"INFO: Article {article['article_number']:03d} -> Page {page_mapping[i]} ({article_pages} pages)")

        return page_mapping

    def create_preliminary_toc(self):
        """Create preliminary TOC to measure its page count"""
        # This method is now replaced by measure_toc_pages()
        toc_pages, toc_docx_path = self.measure_toc_pages(use_real_page_numbers=False)
        return toc_pages

    def regenerate_final_toc(self, page_mapping):
        """Regenerate TOC with correct page numbers and verify page count"""
        print("INFO: Regenerating TOC with correct page numbers...")

        # Create final TOC and measure
        final_toc_pages, toc_docx_path = self.measure_toc_pages(use_real_page_numbers=True, page_mapping=page_mapping)

        return toc_docx_path, final_toc_pages

    def calculate_final_pages(self, toc_pages):
        """Calculate exact page numbers for each article - LEGACY METHOD"""
        # This method is replaced by calculate_exact_pages()
        return self.calculate_exact_pages(2, toc_pages)  # Assuming 2 title pages

    def merge_all_documents(self, title_docx_path, toc_docx_path, output_filename="final_complete.docx"):
        """Merge all documents using advanced XML merger"""
        print("INFO: Merging all documents with XML-level preservation...")

        try:
            # Prepare list of all documents to merge in order
            documents_to_merge = []

            # 1. Title and blank pages
            documents_to_merge.append(title_docx_path)
            print(f"INFO: Queued title+blank: {Path(title_docx_path).name}")

            # 2. TOC
            documents_to_merge.append(toc_docx_path)
            print(f"INFO: Queued TOC: {Path(toc_docx_path).name}")

            # 3. All individual articles in order
            for article in self.page_counts_data:
                article_file = self.articles_dir / f"article_{article['article_number']:03d}.docx"
                if article_file.exists():
                    documents_to_merge.append(article_file)
                    print(f"INFO: Queued article {article['article_number']:03d}: {article['title'][:30]}...")
                else:
                    print(f"ERROR: Article file not found: {article_file}")

            # Create base document (empty)
            base_doc = Document()
            styles_manager = create_styles_for_document(base_doc, verbose=False)
            base_docx_path = self.output_dir / "00_base_empty.docx"
            base_doc.save(str(base_docx_path))

            # Use advanced XML merger
            final_docx_path = self.output_dir / output_filename
            success = self.merger.merge_docx_xml_advanced(
                target_docx_path=base_docx_path,
                source_docx_paths=documents_to_merge,
                output_path=final_docx_path
            )

            if not success:
                print("ERROR: XML merger failed")
                return None, None

            print(f"INFO: Final DOCX saved: {final_docx_path}")

            # Convert final document to PDF
            final_pdf_path = self.convert_docx_to_pdf(final_docx_path)
            if final_pdf_path:
                print(f"INFO: Final PDF saved: {final_pdf_path}")

                # Count total pages
                total_pages = self.count_pdf_pages(final_pdf_path)
                print(f"INFO: Final document has {total_pages} pages total")

            # Cleanup base document
            try:
                base_docx_path.unlink()
            except:
                pass

            return final_docx_path, final_pdf_path

        except Exception as e:
            print(f"ERROR: Failed to merge documents: {e}")
            return None, None

    def convert_docx_to_pdf(self, docx_path):
        """Convert DOCX to PDF using LibreOffice"""
        try:
            docx_path = Path(docx_path)
            pdf_path = docx_path.with_suffix('.pdf')

            cmd = [
                'libreoffice',
                '--headless',
                '--convert-to', 'pdf',
                str(docx_path.name)
            ]

            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=120,
                cwd=str(docx_path.parent)
            )

            if pdf_path.exists():
                return str(pdf_path)
            else:
                print(f"ERROR: PDF not created for {docx_path.name}")
                return None

        except subprocess.TimeoutExpired:
            print(f"ERROR: LibreOffice timeout for {docx_path}")
            return None
        except Exception as e:
            print(f"ERROR: LibreOffice failed: {e}")
            return None

    def count_pdf_pages(self, pdf_path):
        """Count pages in a PDF file"""
        try:
            reader = PdfReader(pdf_path)
            return len(reader.pages)
        except Exception as e:
            print(f"ERROR: Could not count pages in {pdf_path}: {e}")
            return 1

    def run_assembly_process(self, output_filename="final_complete.docx"):
        """Run the complete assembly process with accurate measurements"""
        print("FINALS - DOCUMENT ASSEMBLY PROCESS")
        print("=" * 60)

        try:
            # Step 1: Load page counts
            if not self.load_page_counts_json():
                return False

            # Step 2: Validate prerequisites
            if not self.validate_prerequisites():
                return False

            # Step 3: Measure title+blank pages (REAL measurement)
            title_pages, title_docx_path = self.measure_title_blank_pages()

            # Step 4: Create preliminary TOC and measure pages (REAL measurement)
            preliminary_toc_pages, _ = self.measure_toc_pages(use_real_page_numbers=False)

            # Step 5: Calculate initial page mapping
            initial_page_mapping = self.calculate_exact_pages(title_pages, preliminary_toc_pages)

            # Step 6: Regenerate TOC with real page numbers and measure again
            final_toc_docx_path, final_toc_pages = self.regenerate_final_toc(initial_page_mapping)

            # Step 7: Check if TOC pages changed (due to longer page numbers)
            if final_toc_pages != preliminary_toc_pages:
                print(f"WARNING: TOC pages changed! Preliminary: {preliminary_toc_pages}, Final: {final_toc_pages}")
                print("INFO: Recalculating page mapping with correct TOC size...")

                # Recalculate with correct TOC pages
                final_page_mapping = self.calculate_exact_pages(title_pages, final_toc_pages)

                # Regenerate TOC one more time with corrected page numbers
                print("INFO: Regenerating TOC with corrected page numbers...")
                final_toc_docx_path, verification_toc_pages = self.regenerate_final_toc(final_page_mapping)

                if verification_toc_pages != final_toc_pages:
                    print(f"ERROR: TOC pages still unstable! {final_toc_pages} -> {verification_toc_pages}")
                    print("WARNING: Proceeding with last measurement...")
                    final_toc_pages = verification_toc_pages
                else:
                    print(f"INFO: TOC pages stabilized at {final_toc_pages}")
            else:
                print(f"INFO: TOC pages consistent: {final_toc_pages}")
                final_page_mapping = initial_page_mapping

            # Step 8: Merge all documents using XML merger
            final_docx_path, final_pdf_path = self.merge_all_documents(
                title_docx_path, final_toc_docx_path, output_filename
            )

            if final_docx_path is None:
                return False

            # Step 9: Final verification
            if final_pdf_path:
                actual_total_pages = self.count_pdf_pages(final_pdf_path)
                expected_total_pages = title_pages + final_toc_pages + sum(article['pages'] for article in self.page_counts_data)

                print(f"\nFINAL VERIFICATION:")
                print(f"  Expected total pages: {expected_total_pages}")
                print(f"  Actual total pages: {actual_total_pages}")

                if actual_total_pages == expected_total_pages:
                    print("✅ PAGE COUNT PERFECT MATCH!")
                else:
                    difference = actual_total_pages - expected_total_pages
                    print(f"⚠️  PAGE COUNT DIFFERENCE: {difference:+d} pages")

            print("\n" + "=" * 60)
            print("ASSEMBLY COMPLETED SUCCESSFULLY!")
            print("=" * 60)
            print(f"✅ Final DOCX: {final_docx_path}")
            print(f"✅ Final PDF: {final_pdf_path}")
            print(f"📊 Total articles: {len(self.page_counts_data)}")
            print(f"📄 Title+Blank pages: {title_pages}")
            print(f"📄 TOC pages: {final_toc_pages}")
            print("=" * 60)

            return True

        except Exception as e:
            print(f"ERROR: Assembly process failed: {e}")
            return False

    def cleanup_temporary_files(self):
        """Clean up temporary assembly files"""
        temp_files = [
            "00_preliminary_toc_measure.docx",
            "00_preliminary_toc_measure.pdf",
            "00_final_toc_measure.docx",
            "00_final_toc_measure.pdf",
            "00_title_blank_measure.docx",
            "00_title_blank_measure.pdf"
        ]

        print("INFO: Cleaning up temporary files...")
        for temp_file in temp_files:
            temp_path = self.output_dir / temp_file
            if temp_path.exists():
                try:
                    # temp_path.unlink()  # Uncomment to enable cleanup
                    print(f"INFO: Would delete {temp_file} (cleanup disabled)")
                except Exception as e:
                    print(f"WARNING: Could not delete {temp_file}: {e}")

        print("INFO: Cleanup completed (files preserved for debugging)")


def main():
    """Main function with command line support"""
    parser = argparse.ArgumentParser(
        description='Finals - Document Assembly Tool\n\n'
                    'Assembles complete DOCX from individual articles with accurate page numbering.\n'
                    'Process:\n'
                    '  1. Loads page counts from articles_page_counts.json\n'
                    '  2. Creates preliminary TOC → measures TOC pages\n'
                    '  3. Calculates exact page numbers for all articles\n'
                    '  4. Regenerates TOC with correct page numbers\n'
                    '  5. Merges: Title + Blank + TOC + All Articles\n'
                    '  6. Post-processes XML to remove duplicate page breaks\n'
                    '  7. Converts final DOCX to PDF\n\n'
                    'Requirements:\n'
                    '  • ./output/articles_page_counts.json\n'
                    '  • ./output/00_single_articles/article_XXX.docx files\n'
                    '  • LibreOffice (for PDF conversion)\n\n'
                    'Example: python Finals.py complete_collection.docx',
        formatter_class=argparse.RawDescriptionHelpFormatter
    )

    parser.add_argument(
        'output_filename',
        nargs='?',
        default='final_complete.docx',
        help='Output filename for final document (default: final_complete.docx)'
    )

    parser.add_argument(
        '--cleanup',
        action='store_true',
        help='Clean up temporary files after assembly'
    )

    parser.add_argument(
        '-v', '--verbose',
        action='store_true',
        help='Enable verbose output'
    )

    args = parser.parse_args()

    try:
        # Create Finals processor
        finals = Finals()

        if args.verbose:
            print(f"Configuration:")
            print(f"  Output filename: {args.output_filename}")
            print(f"  Articles directory: {finals.articles_dir}")
            print(f"  Page counts file: {finals.page_counts_file}")
            print(f"  Cleanup enabled: {args.cleanup}")
            print()

        # Run assembly process
        success = finals.run_assembly_process(args.output_filename)

        if success:
            if args.cleanup:
                finals.cleanup_temporary_files()

            print(f"\n🎉 SUCCESS: Final document assembly completed!")
            print(f"📁 Check ./output/ for your final files")
        else:
            print(f"\n❌ FAILED: Document assembly failed")
            return 1

        return 0

    except KeyboardInterrupt:
        print("\nOperation cancelled by user")
        return 130

    except Exception as e:
        print(f"ERROR: {e}")
        if args.verbose:
            import traceback
            traceback.print_exc()
        return 1


if __name__ == "__main__":
    sys.exit(main())
